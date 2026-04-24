"""Dokumente lesen: PDF, DOCX, TXT, Bilder.

Regeln:
- OCR wird IMMER ausgefuehrt (nicht nur als Fallback) fuer absolute Sicherheit
- Lese-Reihenfolge: links oben → rechts unten (Tesseract --psm 3 --oem 3)
- Alle Seiten werden verarbeitet (kein max_pages-Limit, Standard = 0 = alle)
- Max. Dateigrösse: 2 GB
- PDF: direkte Textextraktion UND OCR, pro Seite wird das laengere Ergebnis genutzt
- DOCX: Paragraphen + Tabellen + Kopf-/Fusszeilen
- Bilder: direkt via Tesseract OCR
- TXT/MD: Encoding-Reihenfolge utf-8 → latin-1 → cp1252 → replace
"""

from __future__ import annotations

import logging
import sys
from pathlib import Path

logger = logging.getLogger(__name__)

def _install_hint_tesseract() -> str:
    if sys.platform == "darwin":
        return "brew install tesseract tesseract-lang poppler"
    elif sys.platform == "win32":
        return (
            "Windows: choco install tesseract  ODER  manuell unter "
            "https://github.com/UB-Mannheim/tesseract/wiki"
        )
    return "sudo apt install tesseract-ocr tesseract-ocr-deu poppler-utils"

def _install_hint_tesseract_only() -> str:
    if sys.platform == "darwin":
        return "brew install tesseract tesseract-lang"
    elif sys.platform == "win32":
        return (
            "Windows: choco install tesseract  ODER  manuell unter "
            "https://github.com/UB-Mannheim/tesseract/wiki"
        )
    return "sudo apt install tesseract-ocr tesseract-ocr-deu"

# Maximale Dateigrösse: 2 GB
MAX_FILE_SIZE_BYTES = 2 * 1024 * 1024 * 1024

# Tesseract-Konfiguration: vollautomatische Segmentierung, bester OCR-Modus
# --psm 3 = Automatische Seitensegmentierung (links→rechts, oben→unten)
# --oem 3 = Bester verfuegbarer OCR-Engine-Modus (LSTM + Legacy)
_TESS_CONFIG = "--psm 3 --oem 3"


def read_text(
    file_path: Path,
    ocr_languages: str = "eng+deu+sqi",
    ocr_dpi: int = 300,
    max_pages: int = 0,
) -> str:
    """Text aus einer Datei extrahieren. OCR wird immer ausgefuehrt.

    Args:
        file_path:     Pfad zur Datei
        ocr_languages: Tesseract-Sprachen (z.B. "eng+deu+sqi")
        ocr_dpi:       Aufloesung fuer PDF→Bild-Konvertierung (Standard: 300)
        max_pages:     Max. Seiten bei PDFs. 0 = alle Seiten (Standard)
    """
    try:
        size = file_path.stat().st_size
        if size > MAX_FILE_SIZE_BYTES:
            logger.warning(
                "Datei zu gross (%.1f MB > 2048 MB), uebersprungen: %s",
                size / 1024 / 1024,
                file_path.name,
            )
            return ""
        if size == 0:
            logger.debug("Leere Datei: %s", file_path.name)
            return ""
    except OSError:
        pass

    ext = file_path.suffix.lower()

    if ext == ".pdf":
        return _read_pdf(file_path, ocr_languages, ocr_dpi, max_pages)
    elif ext == ".docx":
        return _read_docx(file_path)
    elif ext in {".txt", ".md"}:
        return _read_text_file(file_path)
    elif ext in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".webp"}:
        return _read_image(file_path, ocr_languages)
    else:
        logger.warning("Nicht unterstuetztes Format: %s", ext)
        return ""


def _read_pdf(file_path: Path, ocr_languages: str, ocr_dpi: int, max_pages: int) -> str:
    """PDF lesen: direkte Textextraktion + OCR auf jeder Seite.

    Pro Seite wird das laengere Ergebnis (direkt vs. OCR) verwendet.
    Lese-Reihenfolge: links oben → rechts unten (Tesseract --psm 3).
    """
    # ── Schritt 1: Direkte Text-Extraktion (alle Seiten) ────────────────────
    direct_parts: list[str] = []
    total_pages = 0
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(file_path))
        total_pages = len(reader.pages)
        pages = reader.pages if max_pages == 0 else reader.pages[:max_pages]
        for page in pages:
            extracted = page.extract_text() or ""
            direct_parts.append(extracted.strip())
        logger.debug(
            "PDF direkte Extraktion: %d Seiten, %d Zeichen gesamt (%s)",
            len(direct_parts),
            sum(len(p) for p in direct_parts),
            file_path.name,
        )
    except PermissionError:
        logger.error("Keine Leseberechtigung: %s", file_path)
        return ""
    except Exception as exc:
        logger.warning("PDF-Textextraktion fehlgeschlagen fuer %s: %s", file_path.name, exc)

    # ── Schritt 2: OCR auf jeder Seite (immer!) ─────────────────────────────
    ocr_parts: list[str] = []
    try:
        from pdf2image import convert_from_path
        import pytesseract

        convert_kwargs: dict = {"dpi": ocr_dpi, "first_page": 1}
        if max_pages > 0:
            convert_kwargs["last_page"] = max_pages

        # Windows: Poppler-Pfad automatisch suchen (Chocolatey-Standard-Pfad)
        if sys.platform == "win32":
            import shutil as _shutil
            if not _shutil.which("pdftoppm"):
                _win_poppler = Path(r"C:\Program Files\poppler\Library\bin")
                if _win_poppler.exists():
                    convert_kwargs["poppler_path"] = str(_win_poppler)

        page_hint = (
            total_pages if max_pages == 0 else min(max_pages, total_pages)
        ) if total_pages > 0 else "?"
        logger.info(
            "OCR gestartet fuer %s (%s Seiten, %d DPI) ...",
            file_path.name,
            page_hint,
            ocr_dpi,
        )
        images = convert_from_path(str(file_path), **convert_kwargs)
        for idx, img in enumerate(images, start=1):
            ocr_text = pytesseract.image_to_string(
                img,
                lang=ocr_languages,
                config=_TESS_CONFIG,
            )
            ocr_parts.append(ocr_text.strip())
            logger.debug("OCR Seite %d/%d: %d Zeichen", idx, len(images), len(ocr_text))

    except FileNotFoundError:
        logger.error(
            "Tesseract oder Poppler nicht gefunden. Installation: %s",
            _install_hint_tesseract(),
        )
    except Exception as exc:
        logger.warning("OCR fehlgeschlagen fuer %s: %s", file_path.name, exc)

    # ── Schritt 3: Beste Ergebnis pro Seite zusammenfuehren ─────────────────
    final_parts: list[str] = []
    num_pages = max(len(direct_parts), len(ocr_parts))
    for i in range(num_pages):
        direct = direct_parts[i] if i < len(direct_parts) else ""
        ocr    = ocr_parts[i]    if i < len(ocr_parts)    else ""
        # Laengeres Ergebnis bevorzugen (mehr Text = mehr Inhalt)
        final_parts.append(ocr if len(ocr) >= len(direct) else direct)

    result = "\n".join(final_parts).strip()
    logger.info(
        "PDF fertig: %s — %d Seiten, %d Zeichen extrahiert",
        file_path.name,
        num_pages,
        len(result),
    )
    return result


def _read_docx(file_path: Path) -> str:
    """DOCX-Datei lesen: Paragraphen + Tabellen + Kopf-/Fusszeilen.

    Liest Body-Elemente in Dokumentreihenfolge (Paragraphen UND Tabellen
    als direkte Kinder des Body-Elements), um Dopplung zu vermeiden:
    document.paragraphs wuerde auch Absaetze in Tabellenzellen liefern.
    Zusammengefuehrte Zellen werden dedupliziert (via Element-ID).
    Kopf-/Fusszeilen nur wenn nicht mit vorheriger Sektion verknuepft.
    """
    try:
        import docx as _docx
        from docx.text.paragraph import Paragraph as _Para
        from docx.table import Table as _Table

        document = _docx.Document(str(file_path))
        parts: list[str] = []
        seen_cell_ids: set[int] = set()  # Deduplizierung zusammengefuehrter Zellen

        # ── Body in Dokumentreihenfolge (nur Top-Level-Elemente) ─────────────
        for block in document.element.body:
            tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag

            if tag == "p":
                p = _Para(block, document)
                t = p.text.strip()
                if t:
                    parts.append(t)

            elif tag == "tbl":
                table = _Table(block, document)
                for row in table.rows:
                    row_texts: list[str] = []
                    for cell in row.cells:
                        cell_id = id(cell._tc)
                        if cell_id in seen_cell_ids:
                            continue  # Zusammengefuehrte Zelle bereits verarbeitet
                        seen_cell_ids.add(cell_id)
                        ct = cell.text.strip()
                        if ct:
                            row_texts.append(ct)
                    if row_texts:
                        parts.append(" | ".join(row_texts))

        # ── Kopf- und Fusszeilen (nur eindeutige Sektionen) ──────────────────
        for section in document.sections:
            for hf in [section.header, section.footer]:
                if hf is None:
                    continue
                if getattr(hf, "is_linked_to_previous", False):
                    continue  # Inhalt wurde bereits von vorheriger Sektion gelesen
                for p in hf.paragraphs:
                    t = p.text.strip()
                    if t:
                        parts.append(t)

        result = "\n".join(parts).strip()
        logger.info("DOCX fertig: %s — %d Zeichen", file_path.name, len(result))
        return result

    except PermissionError:
        logger.error("Keine Leseberechtigung: %s", file_path)
        return ""
    except Exception as exc:
        logger.warning("DOCX-Lesen fehlgeschlagen fuer %s: %s", file_path.name, exc)
        return ""


def _read_text_file(file_path: Path) -> str:
    """Textdatei lesen -- versucht mehrere Encodings.

    Reihenfolge: utf-8-sig zuerst (entfernt BOM falls vorhanden),
    dann utf-8 (ohne BOM), dann latin-1, cp1252.
    """
    for encoding in ("utf-8-sig", "utf-8", "latin-1", "cp1252"):
        try:
            return file_path.read_text(encoding=encoding).strip()
        except (UnicodeDecodeError, UnicodeError):
            continue
        except PermissionError:
            logger.error("Keine Leseberechtigung: %s", file_path)
            return ""
        except Exception as exc:
            logger.warning("Textdatei-Lesen fehlgeschlagen fuer %s: %s", file_path.name, exc)
            return ""
    # Letzter Ausweg: ersetze unlesbare Zeichen
    try:
        return file_path.read_text(encoding="utf-8", errors="replace").strip()
    except Exception:
        return ""


def _read_image(file_path: Path, ocr_languages: str) -> str:
    """Bild per OCR lesen. Lese-Reihenfolge: links oben → rechts unten.

    Konvertiert RGBA-, Palette- (P), CMYK- und Graustufen+Alpha-Bilder (LA)
    zu RGB, da Tesseract damit am zuverlaessigsten arbeitet.
    """
    try:
        from PIL import Image
        import pytesseract
        with Image.open(file_path) as img:
            # Bildmodus normalisieren: Tesseract arbeitet am besten mit RGB/L
            if img.mode in ("RGBA", "P", "LA", "CMYK"):
                img = img.convert("RGB")
            result = pytesseract.image_to_string(
                img,
                lang=ocr_languages,
                config=_TESS_CONFIG,
            ).strip()
        logger.info("Bild OCR fertig: %s — %d Zeichen", file_path.name, len(result))
        return result
    except FileNotFoundError:
        logger.error(
            "Tesseract nicht gefunden. Installation: %s",
            _install_hint_tesseract_only(),
        )
        return ""
    except ImportError:
        logger.error("Pillow oder pytesseract nicht installiert")
        return ""
    except Exception as exc:
        logger.warning("Bild-OCR fehlgeschlagen fuer %s: %s", file_path.name, exc)
        return ""
